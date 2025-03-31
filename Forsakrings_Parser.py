
import re
import json
from io import BytesIO
from docx import Document
import pandas as pd

BASBELOPP_2025 = 58800

def to_number(varde):
    try:
        if varde is None:
            return 0
        if isinstance(varde, (int, float)):
            return int(varde)
        s = str(varde).lower().replace(" ", "").replace(",", ".").replace("sek", "").replace("kr", "")
        if "basbelopp" in s or "bb" in s:
            return int(float(re.findall(r"(\d+\.?\d*)", s)[0]) * BASBELOPP_2025)
        if "msek" in s or "miljoner" in s:
            return int(float(re.findall(r"(\d+\.?\d*)", s)[0]) * 1_000_000)
        if "k" in s:
            return int(float(re.findall(r"(\d+\.?\d*)", s)[0]) * 1_000)
        digits = ''.join(filter(lambda x: x.isdigit() or x == '.', s))
        return int(float(digits)) if digits else 0
    except:
        return 0

def extract_multiple_amounts(pattern, text):
    return sum([to_number(val) for val in re.findall(pattern, text)])

def poangsatt_villkor(lista):
    df = pd.DataFrame(lista)
    df["Premie"] = df["premie"].apply(to_number)
    df["Självrisk"] = df["självrisk"].apply(to_number)
    df["Egendom"] = df["forsakringsbelopp_egendom"]
    df["Ansvar"] = df["forsakringsbelopp_ansvar"]
    df["Avbrott"] = df["forsakringsbelopp_avbrott"]
    max_p, max_s, max_e, max_a, max_v = df["Premie"].max(), df["Självrisk"].max(), df["Egendom"].max(), df["Ansvar"].max(), df["Avbrott"].max()
    maxify = lambda v, m: round((v / m * 10) if m else 0, 2)
    minify = lambda v, m: round((1 - v / m) * 10 if m else 0, 2)
    df["Poäng_premie"] = df["Premie"].apply(lambda x: minify(x, max_p))
    df["Poäng_självrisk"] = df["Självrisk"].apply(lambda x: minify(x, max_s))
    df["Poäng_egendom"] = df["Egendom"].apply(lambda x: maxify(x, max_e))
    df["Poäng_ansvar"] = df["Ansvar"].apply(lambda x: maxify(x, max_a))
    df["Poäng_avbrott"] = df["Avbrott"].apply(lambda x: maxify(x, max_v))
    df["Totalpoäng"] = df[["Poäng_premie", "Poäng_självrisk", "Poäng_egendom", "Poäng_ansvar", "Poäng_avbrott"]].mean(axis=1).round(2)
    df_sorted = df.sort_values(by="Totalpoäng", ascending=False).reset_index(drop=True)
    benchmark = {
        "Snittpremie": int(df["Premie"].mean()),
        "Snittsjälvrisk": int(df["Självrisk"].mean()),
        "Snittpoäng": round(df["Totalpoäng"].mean(), 2)
    }
    return df_sorted, benchmark

def fargstil(value):
    if value >= 8:
        return 'background-color: #c4f5c2'
    elif value >= 6:
        return 'background-color: #fff4a3'
    elif value >= 4:
        return 'background-color: #ffd2a3'
    else:
        return 'background-color: #ffb6b6'

def render_resultat(df, benchmark, st):
    st.subheader("📊 Sammanställning & poängsättning")
    st.dataframe(df.style.applymap(fargstil, subset=["Totalpoäng"]))
    st.subheader("📉 Benchmarking")
    st.markdown(f"**Snittpremie:** {benchmark['Snittpremie']:,} kr  ")
    st.markdown(f"**Snittsjälvrisk:** {benchmark['Snittsjälvrisk']:,} kr  ")
    st.markdown(f"**Snittpoäng:** {benchmark['Snittpoäng']:.2f}")
    st.subheader("⬇️ Export")
    doc = Document()
    doc.add_heading("Försäkringsjämförelse", level=1)
    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = col
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            row_cells[i].text = str(row[col])
    word_buf = BytesIO()
    doc.save(word_buf)
    word_buf.seek(0)
    json_buf = BytesIO()
    json_buf.write(json.dumps(df.to_dict(orient="records"), indent=2, ensure_ascii=False).encode("utf-8"))
    json_buf.seek(0)
    st.download_button("📄 Ladda ner Word-dokument", data=word_buf, file_name="forsakringsjamforelse.docx")
    st.download_button("🧾 Exportera som JSON", data=json_buf, file_name="forsakringsdata.json")
