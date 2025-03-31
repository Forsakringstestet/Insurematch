import streamlit as st
import pandas as pd
import re
import json
import pdfplumber
from PyPDF2 import PdfReader
from docx import Document
from io import BytesIO
from datetime import date, timedelta

# === Konstanter ===
BASBELOPP_2025 = 58800

# === Valör-konverterare ===
def to_number(varde):
    try:
        if varde is None:
            return 0
        if isinstance(varde, (int, float)):
            return int(varde)
        s = str(varde).lower()
        s = s.replace(" ", "").replace(",", ".").replace("sek", "").replace("kr", "")
        if "basbelopp" in s or "bb" in s:
            val = float(re.findall(r"(\d+\.?\d*)", s)[0])
            return int(val * BASBELOPP_2025)
        if "msek" in s or "miljoner" in s:
            val = float(re.findall(r"(\d+\.?\d*)", s)[0])
            return int(val * 1_000_000)
        if "k" in s:
            val = float(re.findall(r"(\d+\.?\d*)", s)[0])
            return int(val * 1_000)
        digits = ''.join(filter(lambda x: x.isdigit() or x == '.', s))
        return int(float(digits)) if digits else 0
    except Exception as e:
        st.warning(f"⚠️ Fel vid konvertering till nummer: {varde} → {e}")
        return 0

# === PDF-läsare med fallback ===
def läs_pdf_text(pdf_file):
    try:
        with pdfplumber.open(pdf_file) as pdf:
            return "\n".join([page.extract_text() or "" for page in pdf.pages])
    except:
        reader = PdfReader(pdf_file)
        return "\n".join([page.extract_text() or "" for page in reader.pages if page.extract_text()])
def extract_by_patterns(text, patterns):
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return to_number(match.group(2))
    return 0

def extract_självrisk(text, patterns):
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            value = match.group(2)
            suffix = match.group(3).lower() if match.group(3) else ""
            if "basbelopp" in suffix or "bb" in suffix:
                return int(float(value.replace(",", ".")) * BASBELOPP_2025)
            else:
                return to_number(value)
    return 0

def extrahera_forsakringsgivare(text):
    bolag = ["if", "gjensidige", "trygg-hansa", "moderna", "protector", "svedea", "folksam", "dina", "länsförsäkringar", "lf"]
    for namn in bolag:
        if re.search(rf"\b{namn}\b", text, re.IGNORECASE):
            return namn.capitalize()
    return "Okänt"

def extrahera_forsakringsnummer(text):
    match = re.search(r"(försäkringsnummer|avtalsnummer)[\s:\-]+([A-Z0-9\-\/]{6,})", text, re.IGNORECASE)
    return match.group(2) if match else ""

def extrahera_forsakringstid(text):
    match = re.search(r"(\d{4}-\d{2}-\d{2})\s*(–|till|-)\s*(\d{4}-\d{2}-\d{2})", text)
    return f"{match.group(1)} – {match.group(3)}" if match else ""

def extrahera_karens(text):
    match = re.search(r"(karens|karensdagar)[^\d]{0,10}(\d{1,3})", text, re.IGNORECASE)
    return f"{match.group(2)} dagar" if match else ""

def extrahera_ansvarstid(text):
    match = re.search(r"(ansvarstid|ersättningstid)[^\d]{0,10}(\d{1,3})", text, re.IGNORECASE)
    return f"{match.group(2)} månader" if match else ""

def extrahera_lank(text):
    match = re.search(r"https?://[^\s]+", text)
    return match.group(0) if match else "PDF"
def extrahera_villkor_ur_pdf(text):
    premie_patterns = [
        r"(premie|total(?!.*belopp)|pris totalt|summa att betala|kostnad)[^\d]{0,20}([\d\s]+)",
        r"pris för tiden[^\d]{0,20}([\d\s]+)"
    ]
    premie = extract_by_patterns(text, premie_patterns)

    självrisk_patterns = [
        r"(självrisk)[^\d]{0,15}([\d\s]+(?:\.\d+)?)(\s*(kr|sek|basbelopp|bb)?)"
    ]
    självrisk = extract_självrisk(text, självrisk_patterns)

    egendom_poster = {"byggnad": 0, "maskiner": 0, "varor": 0}
    egendom_patterns = {
        "byggnad": r"(byggnad|verkstad|fastighet)[^\d]{0,20}([\d\s]+)",
        "maskiner": r"(maskiner|inventarier)[^\d]{0,20}([\d\s]+)",
        "varor": r"(varor|lager)[^\d]{0,20}([\d\s]+)"
    }
    for key, pattern in egendom_patterns.items():
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            egendom_poster[key] = to_number(match.group(2))
    försäkringsbelopp_egendom = sum(egendom_poster.values())

    ansvar_poster = {"produkt": 0, "allmänt": 0}
    ansvar_patterns = {
        "produkt": r"(produktansvar)[^\d]{0,20}([\d\s]+)",
        "allmänt": r"(ansvarsförsäkring|verksamhetsansvar)[^\d]{0,20}([\d\s]+)"
    }
    for key, pattern in ansvar_patterns.items():
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            ansvar_poster[key] = to_number(match.group(2))
    försäkringsbelopp_ansvar = sum(ansvar_poster.values())

    avbrott_poster = {"täckningsbidrag": 0, "intäktsbortfall": 0}
    avbrott_patterns = {
        "täckningsbidrag": r"(täckningsbidrag|täcknings.*förlust)[^\d]{0,20}([\d\s]+)",
        "intäktsbortfall": r"(intäktsbortfall|förlorad omsättning)[^\d]{0,20}([\d\s]+)"
    }
    for key, pattern in avbrott_patterns.items():
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            avbrott_poster[key] = to_number(match.group(2))
    försäkringsbelopp_avbrott = sum(avbrott_poster.values())

    return {
        "försäkringsgivare": extrahera_forsakringsgivare(text),
        "premie": premie,
        "självrisk": självrisk,
        "försäkringsbelopp_egendom": försäkringsbelopp_egendom,
        "försäkringsbelopp_ansvar": försäkringsbelopp_ansvar,
        "försäkringsbelopp_avbrott": försäkringsbelopp_avbrott,
        "egendom_byggnad": egendom_poster["byggnad"],
        "egendom_maskiner": egendom_poster["maskiner"],
        "egendom_varor": egendom_poster["varor"],
        "ansvar_allmänt": ansvar_poster["allmänt"],
        "ansvar_produkt": ansvar_poster["produkt"],
        "avbrott_täckningsbidrag": avbrott_poster["täckningsbidrag"],
        "avbrott_intäktsbortfall": avbrott_poster["intäktsbortfall"],
        "karens": extrahera_karens(text),
        "ansvarstid": extrahera_ansvarstid(text),
        "försäkringstid": extrahera_forsakringstid(text),
        "försäkringsnummer": extrahera_forsakringsnummer(text),
        "villkorsreferens": extrahera_lank(text),
        "undantag": ""
    }
def färgschema(value):
    if value >= 8:
        return 'background-color: #b6fcb6'
    elif value >= 6:
        return 'background-color: #f9fcb6'
    elif value >= 4:
        return 'background-color: #fde2b6'
    else:
        return 'background-color: #fcb6b6'
def generera_word_dokument(data):
    doc = Document()
    doc.add_heading("Upphandlingsunderlag – Försäkringsjämförelse", level=1)
    table = doc.add_table(rows=1, cols=len(data[0]))
    hdr_cells = table.rows[0].cells
    for i, key in enumerate(data[0].keys()):
        hdr_cells[i].text = key
    for row in data:
        cells = table.add_row().cells
        for i, key in enumerate(row):
            cells[i].text = str(row[key])
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
def generera_json(data):
    buffer = BytesIO()
    buffer.write(json.dumps(data, indent=2, ensure_ascii=False).encode("utf-8"))
    buffer.seek(0)
    return buffer
def poangsatt_villkor(villkor_list):
    df = pd.DataFrame(villkor_list)

    df["Premie"] = df["premie"].apply(to_number)
    df["Självrisk"] = df["självrisk"].apply(to_number)
    df["Egendom"] = df["försäkringsbelopp_egendom"]
    df["Ansvar"] = df["försäkringsbelopp_ansvar"]
    df["Avbrott"] = df["försäkringsbelopp_avbrott"]

    df["Premie_poäng"] = 1 / (df["Premie"] + 1)
    df["Självrisk_poäng"] = 1 / (df["Självrisk"] + 1)
    df["Egendom_poäng"] = df["Egendom"]
    df["Ansvar_poäng"] = df["Ansvar"]
    df["Avbrott_poäng"] = df["Avbrott"]

    for col in ["Premie_poäng", "Självrisk_poäng", "Egendom_poäng", "Ansvar_poäng", "Avbrott_poäng"]:
        max_val = df[col].max()
        df[col] = df[col] / max_val * 10 if max_val > 0 else 0

    def karens_poäng(k):
        if "1" in str(k): return 0.5
        if "2" in str(k): return 0.2
        if "3" in str(k): return 0
        return -0.5

    def ansvarstid_poäng(a):
        try:
            val = int(re.search(r"\d+", str(a)).group())
            if val >= 12: return 0.5
            if val >= 6: return 0.2
        except:
            return 0
        return 0

    df["Bonus_karens"] = df["karens"].apply(karens_poäng)
    df["Bonus_ansvarstid"] = df["ansvarstid"].apply(ansvarstid_poäng)

    df["Totalpoäng"] = (
        df["Premie_poäng"] * 0.20 +
        df["Självrisk_poäng"] * 0.15 +
        df["Egendom_poäng"] * 0.25 +
        df["Ansvar_poäng"] * 0.25 +
        df["Avbrott_poäng"] * 0.15 +
        df["Bonus_karens"] +
        df["Bonus_ansvarstid"]
    ).round(2)

    df.rename(columns={
        "försäkringsgivare": "Försäkringsgivare",
        "undantag": "Undantag",
        "villkorsreferens": "Källa",
        "försäkringstid": "Försäkringstid",
        "försäkringsnummer": "Försäkringsnummer"
    }, inplace=True)

    return df[[
        "Försäkringsgivare", "Premie", "Självrisk", "Egendom", "Ansvar", "Avbrott",
        "Försäkringstid", "Försäkringsnummer", "karens", "ansvarstid", "Undantag", "Källa", "Totalpoäng"
    ]]
if __name__ == "__main__":
    st.set_page_config(page_title="Försäkringsguide", page_icon="🛡️", layout="centered")
    st.title("🛡️ Försäkringsguide & Jämförelse")

    uploaded_pdfs = st.file_uploader("📂 Ladda upp en eller flera PDF:er", type="pdf", accept_multiple_files=True)
    påminnelse_datum = st.date_input("🔔 Vill du få en påminnelse innan förnyelse?", value=date.today() + timedelta(days=300))

    if uploaded_pdfs:
        visa_rådata = st.checkbox("🧾 Visa extraherade rådata")
        villkorslista = []
        st.markdown("### 📄 Offertanalys")

        for i, pdf in enumerate(uploaded_pdfs):
            text = läs_pdf_text(pdf)
            st.markdown(f"#### 📑 Fil {i+1}: {pdf.name}")
            st.text_area("📃 PDF-text (förhandsvisning)", value=text[:2000], height=200)

            extrakt = extrahera_villkor_ur_pdf(text)
            villkorslista.append(extrakt)

            if visa_rådata:
                st.json(extrakt)

            with st.expander("📂 Visa delbelopp"):
                st.markdown(f"""
                **Egendom**  
                - 🧰 Maskiner: `{extrakt.get("egendom_maskiner", 0):,} kr`  
                - 🏗️ Byggnad: `{extrakt.get("egendom_byggnad", 0):,} kr`  
                - 📦 Varor: `{extrakt.get("egendom_varor", 0):,} kr`  

                **Ansvar**  
                - ⚖️ Produktansvar: `{extrakt.get("ansvar_produkt", 0):,} kr`  
                - 📜 Allmänt ansvar: `{extrakt.get("ansvar_allmänt", 0):,} kr`

                **Avbrott**  
                - 💸 Täckningsbidrag: `{extrakt.get("avbrott_täckningsbidrag", 0):,} kr`  
                - 📉 Intäktsbortfall: `{extrakt.get("avbrott_intäktsbortfall", 0):,} kr`
                """)

            saknade = [k for k, v in extrakt.items() if to_number(v) == 0 and k not in ["undantag", "villkorsreferens", "karens", "ansvarstid"]]
            if saknade:
                st.warning(f"⚠️ Saknade värden i {pdf.name}: {', '.join(saknade)}")

            st.markdown("---")
        if villkorslista:
            df = poangsatt_villkor(villkorslista)

            st.subheader("📊 Sammanställning & poängsättning")
            st.dataframe(df.style.applymap(färgschema, subset=["Totalpoäng"]))

            st.markdown("### 📉 Benchmarking")
            st.markdown(f"""
                **Snittpremie:** {df['Premie'].mean():,.0f} kr  
                **Snittsjälvrisk:** {df['Självrisk'].mean():,.0f} kr  
                **Snittpoäng:** {df['Totalpoäng'].mean():.2f}
            """)

            st.download_button(
                "⬇️ Ladda ner sammanställning (Word)",
                data=generera_word_dokument(df.to_dict(orient="records")),
                file_name="jamforelse_upphandling.docx"
            )

            st.download_button(
                "💾 Exportera som JSON",
                data=generera_json(villkorslista),
                file_name="jamforelse_data.json"
            )

            st.success(f"🔔 Påminnelse sparad: Lägg in {påminnelse_datum} i din kalender")
